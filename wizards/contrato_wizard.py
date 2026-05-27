# -*- coding: utf-8 -*-
from odoo import models, fields, api
from odoo.exceptions import UserError
import base64
import copy
import subprocess
import tempfile
import os
import re as _re


class ContratoWizard(models.TransientModel):
    _name        = 'reclutamiento__kuale.contrato_wizard'
    _description = 'Generar Contrato de Trabajo'

    applicant_id       = fields.Many2one('hr.applicant', required=True)
    formatos_wizard_id = fields.Many2one('reclutamiento__kuale.formatos_wizard')

    contrato_id = fields.Many2one(
        'hr.applicant.doc.format',
        string='Tipo de contrato',
        domain="[('docx_file', '!=', False)]",
        required=True,
    )

    is_determinado = fields.Boolean(
        string='Es determinado',
        compute='_compute_is_determinado',
    )

    @api.depends('contrato_id')
    def _compute_is_determinado(self):
        for rec in self:
            name = (rec.contrato_id.name or '').lower()
            rec.is_determinado = 'determinado' in name and 'indeterminado' not in name

    @api.model
    def default_get(self, fields_list):
        defaults = super().default_get(fields_list)
        if 'contrato_id' in fields_list and not defaults.get('contrato_id'):
            contrato = self.env['hr.applicant.doc.format'].search(
                [('docx_file', '!=', False), ('active', '=', True)],
                limit=1, order='id asc',
            )
            if contrato:
                defaults['contrato_id'] = contrato.id
        return defaults
    
    contract_duration = fields.Selection([
        ('30', '30 días'),
        ('60', '60 días'),
        ('90', '90 días'),
    ], string='Duración del contrato')

    contract_end_date = fields.Date(
        string='Fecha fin de contrato',
        compute='_compute_contract_end_date',
        store=False,
    )

    @api.depends('contract_duration', 'applicant_id')
    def _compute_contract_end_date(self):
        from datetime import date, timedelta
        for rec in self:
            if rec.contract_duration:
                # Usar fecha de inicio del applicant, si no hoy
                base = (rec.applicant_id.contract_start_date 
                        if rec.applicant_id and rec.applicant_id.contract_start_date 
                        else date.today())
                rec.contract_end_date = base + timedelta(days=int(rec.contract_duration))
            else:
                rec.contract_end_date = False

    def _get_docx_bytes(self):
        fmt = self.contrato_id
        if not fmt or not fmt.docx_file:
            raise UserError(
                'El contrato seleccionado no tiene archivo .docx. '
                'Súbelo en Catálogo → Formatos de contratación.'
            )
        return base64.b64decode(fmt.docx_file)

    def _build_replacements(self, applicant):
        a = applicant

        full_name = ' '.join(filter(None, [
            a.partner_name or '',
            a.last_name or '',
            a.last_name2 or '',
        ])).strip()

        rep_legal = ''
        if a.company_id and hasattr(a.company_id, 'legal_representative') and a.company_id.legal_representative:
            rep_legal = a.company_id.legal_representative.name or ''

        _company = a.company_id
        if _company and _company.parent_id:
            _company = _company.parent_id

        company_rfc   = getattr(_company, 'rfc', '') or getattr(_company, 'vat', '') or ''
        company_name  = _company.name if _company else ''
        company_city  = (_company.city or '') if _company else ''
        company_state = (_company.state_id.name if _company and _company.state_id else '')
        company_addr  = ''
        if _company:
            parts = list(filter(None, [
                _company.street or '',
                _company.city or '',
                _company.state_id.name if _company.state_id else '',
                'C.P. ' + (_company.zip or '') if _company.zip else '',
            ]))
            company_addr = ', '.join(parts)

        # Salario: applicant primero, luego job
        salary = a.salary_proposed or 0
        if not salary and a.job_id:
            salary = a.job_id.net_base_salary or 0
        salary_fmt = '${:,.2f}'.format(salary)

        from datetime import date
        today = date.today()
        start_date = getattr(applicant, 'contract_start_date', None) or today
        months_es = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }

        # Descripción de contrato (campo exclusivo, texto plano)
        job_desc = (a.job_id.contract_description or '') if a.job_id else ''

        # Horas semanales del horario de trabajo del puesto
        horas_semana = ''
        if a.job_id and hasattr(a.job_id, 'working_schedule_id') and a.job_id.working_schedule_id:
            schedule = a.job_id.working_schedule_id
            try:
                import re as _re2
                match = _re2.search(r'(\d+(?:\.\d+)?)\s*h', schedule.name or '', _re2.IGNORECASE)
                if match:
                    h = float(match.group(1))
                    horas_semana = str(int(h)) if h == int(h) else str(h)
                else:
                    # Fallback: sumar líneas si el nombre no tiene número
                    total_hours = sum((line.hour_to - line.hour_from) for line in schedule.attendance_ids)
                    horas_semana = str(int(total_hours)) if total_hours == int(total_hours) else str(round(total_hours, 1))
            except Exception:
                horas_semana = ''

        # Salary text: convertir el salario a texto en español
        try:
            from num2words import num2words as _n2w
            _entero   = int(salary)
            _centavos = round((salary - _entero) * 100)
            salary_text = _n2w(_entero, lang='es').upper()
            salary_text += f' PESOS {_centavos:02d}/100'
        except Exception:
            salary_text = ''

        try:
            nationality_map = dict(a._fields['nationality'].selection)
        except Exception:
            nationality_map = {}
        try:
            marital_map = dict(a._fields['marital_status'].selection)
        except Exception:
            marital_map = {}

        nationality = nationality_map.get(a.nationality, '') if a.nationality else ''
        marital     = marital_map.get(a.marital_status, '') if a.marital_status else ''
        curp        = getattr(a, 'curp', '') or ''
        nss         = getattr(a, 'social_security_number', '') or ''
        edad        = str(a.age) if getattr(a, 'age', None) else ''
        direccion   = ' '.join(filter(None, [
            getattr(a, 'current_address', '') or '',
            getattr(a, 'exterior_number', '') or '',
            getattr(a, 'colony', '') or '',
            getattr(a, 'municipality', '') or '',
            getattr(a, 'state', '') or '',
            'C.P. ' + (getattr(a, 'postal_code', '') or '') if getattr(a, 'postal_code', '') else '',
        ]))
        telefono = a.partner_mobile or a.partner_phone or ''

        fecha_fin = ''
        if self.contract_end_date:
            fecha_fin = self.contract_end_date.strftime('%d/%m/%Y')

        def U(val):
            return (val or '').upper()

        return {
            '--rep_legal--':      U(rep_legal),
            '--company_name--':   U(company_name),
            '--company_rfc--':    U(company_rfc),
            '--company_city--':   U(company_city),
            '--company_state--':  U(company_state),
            '--company_addr--':   U(company_addr),
            '--employee_name--':  U(full_name),
            '--salary_fmt--':     salary_fmt,
            '--salary_text--':    U(salary_text),
            '--job_desc--':       U(job_desc),
            '--Horas a la semana--': horas_semana,
            '--nationality--':    U(nationality),
            '--marital--':        U(marital),
            '--curp--':           U(curp),
            '--nss--':            U(nss),
            '--edad--':           U(edad),
            '--direccion--':      U(direccion),
            '--telefono--':       U(telefono),
            '--dia--':            str(start_date.day),
            '--mes--':            months_es[start_date.month],
            '--anio--':           str(start_date.year),
            '--fecha_fin--':      U(fecha_fin),
            # ── Generales empleado (blancos) ──────────────────────────────────
            'NACIONALIDAD: _____________________________': 'NACIONALIDAD: ' + U(nationality),
            'ESTADO CIVIL: _____________________________': 'ESTADO CIVIL: ' + U(marital),
            'C.U.R.P.: __________________________________': 'C.U.R.P.: ' + U(curp),
            'NUMERO DE AFILIACION AL IMSS: _____________________________': 'NUMERO DE AFILIACION AL IMSS: ' + U(nss),
            'EDAD: __________________': 'EDAD: ' + U(edad),
            'DIRECCION: ______________________________________________________________': 'DIRECCION: ' + U(direccion),
            'TELEFONO: _____________________': 'TELEFONO: ' + U(telefono),
            '                    ______________________________________________________________': '',
        }

    # ──────────────────────────────────────────────────────────────────────────
    # REEMPLAZO DE VARIABLES — sin modificar la estructura del documento
    # ──────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _rpr_key(run):
        from lxml import etree
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        rpr = run._r.find(f'{{{ns}}}rPr')
        return etree.tostring(rpr, encoding='unicode') if rpr is not None else ''

    def _replace_text_in_paragraph(self, paragraph, replacements):
        if not paragraph.runs:
            return

        full_text = ''.join(r.text for r in paragraph.runs)
        if not any(old in full_text for old in replacements):
            return

        # 1. Agrupar runs consecutivos con mismo rPr
        groups = []
        current_group = []
        current_key = None
        for run in paragraph.runs:
            key = self._rpr_key(run)
            if key == current_key:
                current_group.append(run)
            else:
                if current_group:
                    groups.append(current_group)
                current_group = [run]
                current_key = key
        if current_group:
            groups.append(current_group)

        # 2. Reemplazar dentro de cada grupo (preserva el rPr del grupo)
        for group in groups:
            group_text = ''.join(r.text for r in group)
            new_text = group_text
            for old, new in replacements.items():
                if old in new_text:
                    new_text = new_text.replace(old, new)
            if new_text != group_text:
                group[0].text = new_text
                for run in group[1:]:
                    run.text = ''

        # 3. Fallback: placeholder que cruza grupos con diferente rPr
        full_after = ''.join(r.text for r in paragraph.runs)
        if not any(old in full_after for old in replacements):
            return

        new_full = full_after
        for old, new in replacements.items():
            if old in new_full:
                new_full = new_full.replace(old, new)
        if new_full == full_after:
            return

        chars = list(new_full)
        pos = 0
        runs = paragraph.runs
        for i, run in enumerate(runs):
            orig_len = len(run.text)
            if i < len(runs) - 1:
                run.text = ''.join(chars[pos:pos + orig_len])
                pos += orig_len
            else:
                run.text = ''.join(chars[pos:])

    def _fill_beneficiaries_table(self, table, applicant):
        """
        Rellena la tabla de beneficiarios usando las filas existentes de la
        plantilla. NO agrega ni elimina filas — solo escribe los datos en las
        filas vacías disponibles (filas 1 en adelante, después del header).
        """
        beneficiaries = getattr(applicant, 'beneficiaries_ids', None)
        data_rows = table.rows[1:]

        if not data_rows:
            return

        rel_map = {}
        if beneficiaries:
            try:
                rel_map = dict(beneficiaries[0]._fields['beneficiary_relationship'].selection)
            except Exception:
                pass

        from docx.oxml.ns import qn

        def set_cell_text(cell, text):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = text
                    text = ''
                    return
                if text:
                    para.add_run(text)
                    return

        for row_idx, row in enumerate(data_rows):
            if beneficiaries and row_idx < len(beneficiaries):
                b = beneficiaries[row_idx]
                rel_label = (rel_map.get(b.beneficiary_relationship, '')
                             or getattr(b, 'other_relationship', '') or '')
                values = [
                    b.beneficiary_name or '',
                    rel_label,
                    str(b.beneficiary_percentage or '') + '%',
                ]
            else:
                values = ['', '', '']

            for cell, value in zip(row.cells, values):
                set_cell_text(cell, value)

    # ──────────────────────────────────────────────────────────────────────────
    # HELPER: obtener bytes del anexo (PDF listo para fusionar)
    # ──────────────────────────────────────────────────────────────────────────

    def _get_annex_pdf_bytes(self, applicant, lo_bin):
        """
        Busca el archivo de anexo en el job del applicant (o en el padre si es
        un puesto hijo). Si el archivo es Word lo convierte a PDF con LibreOffice.
        Devuelve bytes del PDF o None si no hay anexo.
        """
        _job = applicant.job_id
        annex_b64  = _job.annex_file    if _job else False
        annex_name = (_job.annex_filename or '') if _job else ''

        # Subir al puesto padre si el hijo no tiene el anexo
        if not annex_b64 and _job and _job.company_id and _job.company_id.parent_id:
            parent_job = self.env['hr.job'].sudo().search([
                ('name', '=', _job.name),
                ('is_Parent_Job', '=', True),
            ], limit=1)
            if parent_job:
                annex_b64  = parent_job.annex_file
                annex_name = parent_job.annex_filename or ''

        if not annex_b64:
            return None

        annex_bytes = base64.b64decode(annex_b64)

        # Si es Word, convertir a PDF primero
        if annex_name.lower().endswith(('.docx', '.doc')):
            suffix = '.docx' if annex_name.lower().endswith('.docx') else '.doc'
            annex_tmp     = None
            annex_pdf_tmp = None
            try:
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                    f.write(annex_bytes)
                    annex_tmp = f.name
                annex_pdf_tmp = annex_tmp.rsplit('.', 1)[0] + '.pdf'
                subprocess.run(
                    [lo_bin, '--headless', '--convert-to', 'pdf',
                     '--outdir', os.path.dirname(annex_tmp), annex_tmp],
                    capture_output=True,
                    timeout=120,
                    **({'creationflags': 0x08000000} if os.name == 'nt' else {}),
                )
                if os.path.exists(annex_pdf_tmp) and os.path.getsize(annex_pdf_tmp) > 0:
                    with open(annex_pdf_tmp, 'rb') as f:
                        return f.read()
                return None
            except Exception:
                return None
            finally:
                if annex_tmp and os.path.exists(annex_tmp):
                    os.unlink(annex_tmp)
                if annex_pdf_tmp and os.path.exists(annex_pdf_tmp):
                    os.unlink(annex_pdf_tmp)

        # Asumir PDF directo
        return annex_bytes

    # ──────────────────────────────────────────────────────────────────────────
    # HELPER: fusionar dos PDFs (bytes + bytes → bytes)
    # ──────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _merge_pdfs(main_bytes, annex_bytes):
        import io as _io
        try:
            from pypdf import PdfWriter, PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfWriter, PdfReader
            except ImportError:
                return main_bytes  # sin librería → devolver contrato sin anexo

        writer = PdfWriter()
        for page in PdfReader(_io.BytesIO(main_bytes)).pages:
            writer.add_page(page)
        for page in PdfReader(_io.BytesIO(annex_bytes)).pages:
            writer.add_page(page)
        out = _io.BytesIO()
        writer.write(out)
        return out.getvalue()

    # ──────────────────────────────────────────────────────────────────────────
    # ACCIÓN PRINCIPAL
    # ──────────────────────────────────────────────────────────────────────────

    def action_generate(self):
        self.ensure_one()

        if not self.contrato_id:
            raise UserError('Selecciona un tipo de contrato.')

        if self.is_determinado and not self.contract_duration:
            raise UserError(
                'Este contrato es por tiempo determinado. '
                'Por favor selecciona la duración.'
            )

        applicant = self.applicant_id

        if not applicant.actual_date:
            applicant.sudo().write({'actual_date': fields.Date.today()})

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise UserError('Instala python-docx: pip install python-docx --break-system-packages')

        import io
        docx_bytes = self._get_docx_bytes()
        doc = DocxDocument(io.BytesIO(docx_bytes))
        replacements = self._build_replacements(applicant)

        # Reemplazar variables en párrafos normales
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)

        # Reemplazar variables en cuadros de texto (txbxContent)
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph as _Paragraph
        for txbx in doc.element.body.findall('.//' + _qn('w:txbxContent')):
            for p_el in txbx.findall(_qn('w:p')):
                self._replace_text_in_paragraph(_Paragraph(p_el, doc), replacements)

        # Procesar tablas
        for table in doc.tables:
            header_text = ''.join(c.text for c in table.rows[0].cells).lower()
            if 'nombre' in header_text and ('parentesco' in header_text or 'porcentaje' in header_text):
                self._fill_beneficiaries_table(table, applicant)
            else:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_text_in_paragraph(p, replacements)

        # ── Convertir a PDF via LibreOffice ───────────────────────────────────
        docx_file = None
        pdf_file  = None
        lo_bin    = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                doc.save(f.name)
                docx_file = f.name

            pdf_file = docx_file.replace('.docx', '.pdf')

            for path in [
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                '/usr/bin/soffice',
                '/usr/bin/libreoffice',
                '/usr/local/bin/soffice',
            ]:
                if os.path.exists(path):
                    lo_bin = path
                    break

            if not lo_bin:
                try:
                    from odoo.tools.misc import find_in_path
                    lo_bin = find_in_path('soffice') or find_in_path('libreoffice')
                except Exception:
                    pass

            if not lo_bin:
                raise UserError(
                    'No se encontró LibreOffice. '
                    'Instálalo en el servidor para convertir contratos a PDF.'
                )

            proc = subprocess.run(
                [lo_bin, '--headless', '--convert-to', 'pdf',
                 '--outdir', os.path.dirname(docx_file), docx_file],
                capture_output=True,
                timeout=120,
                **({'creationflags': 0x08000000} if os.name == 'nt' else {}),
            )

            if not os.path.exists(pdf_file) or os.path.getsize(pdf_file) == 0:
                raise UserError(
                    'LibreOffice no generó el PDF.\n'
                    'Error: %s' % (proc.stderr.decode('utf-8', errors='ignore') or 'desconocido')
                )

            with open(pdf_file, 'rb') as f:
                pdf_content = f.read()

        except UserError:
            raise
        except Exception as e:
            raise UserError('Error generando el contrato: %s' % str(e))
        finally:
            if docx_file and os.path.exists(docx_file):
                os.unlink(docx_file)
            if pdf_file and os.path.exists(pdf_file):
                os.unlink(pdf_file)

        # ── Fusionar Anexo si el tipo de contrato lo incluye ─────────────────
        contrato_name = (self.contrato_id.name or '').lower()
        if 'anexo' in contrato_name and lo_bin:
            try:
                annex_pdf_bytes = self._get_annex_pdf_bytes(applicant, lo_bin)
                if annex_pdf_bytes:
                    pdf_content = self._merge_pdfs(pdf_content, annex_pdf_bytes)
            except Exception:
                # Si falla el anexo, el contrato se entrega sin él (no bloquear)
                pass

        # ── Crear adjunto y devolver URL de descarga ──────────────────────────
        filename = 'Contrato_%s.pdf' % (
            applicant.partner_name or 'candidato'
        ).replace(' ', '_')

        att = self.env['ir.attachment'].sudo().create({
            'name':      filename,
            'type':      'binary',
            'datas':     base64.b64encode(pdf_content),
            'res_model': self._name,
            'res_id':    self.id,
            'mimetype':  'application/octet-stream',
        })

        download_url = '/web/content/%s?download=true' % att.id
        formatos_wizard_id = self.formatos_wizard_id.id if self.formatos_wizard_id else False

        return {
            'type':   'ir.actions.client',
            'tag':    'kuale_generar_contrato',
            'params': {
                'download_url':       download_url,
                'formatos_wizard_id': formatos_wizard_id,
            },
        }

    def action_cancel(self):
        self.ensure_one()
        if self.formatos_wizard_id:
            return {
                'type':      'ir.actions.act_window',
                'res_model': 'reclutamiento__kuale.formatos_wizard',
                'view_mode': 'form',
                'res_id':    self.formatos_wizard_id.id,
                'target':    'new',
                'context':   {'dialog_size': 'medium'},
            }
        return {'type': 'ir.actions.act_window_close'}