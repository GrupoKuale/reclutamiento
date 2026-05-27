from markupsafe import Markup
from odoo import models, fields, api
import logging
_logger = logging.getLogger(__name__)


class HojaMembretada(models.Model):
    _name = 'reclutamiento__kuale.hoja_membretada'
    _description = 'Hoja Membretada'

    name = fields.Char(string='Nombre', required=True)
    filename = fields.Char(string='Archivo', required=True)
    image_url = fields.Char(
        string='URL de imagen',
        compute='_compute_image_url',
        store=False
    )
    active = fields.Boolean(default=True)

    @api.depends('filename')
    def _compute_image_url(self):
        for rec in self:
            rec.image_url = (
                '/reclutamiento__kuale/static/src/hojas_membretadas/%s' % rec.filename
            )


class FormatosPreviewWizard(models.TransientModel):
    _name = 'reclutamiento__kuale.formatos_preview_wizard'
    _description = 'Vista previa de formato con hoja membretada'

    applicant_id = fields.Many2one('hr.applicant', string='Candidato', required=True)
    format_id = fields.Many2one(
        'reclutamiento__kuale.format_employee',
        string='Formato', required=True
    )
    hoja_id = fields.Many2one(
        'reclutamiento__kuale.hoja_membretada',
        string='Hoja membretada',
        required=True
    )
    body_html = fields.Html(
        string='Contenido del documento',
        sanitize=False
    )
    hoja_url = fields.Char(
        string='URL hoja membretada',
        compute='_compute_hoja_url',
        store=False
    )

    @api.depends('hoja_id')
    def _compute_hoja_url(self):
        for rec in self:
            rec.hoja_url = rec.hoja_id.image_url if rec.hoja_id else ''

    @api.onchange('hoja_id')
    def _onchange_hoja_id(self):
        pass

    @api.model
    def create_for_format(self, applicant_id, format_id):
        applicant = self.env['hr.applicant'].browse(applicant_id).sudo()
        formato = self.env['reclutamiento__kuale.format_employee'].browse(format_id).sudo()
        body = self._render_format_body(applicant, formato)
        hoja = self.env['reclutamiento__kuale.hoja_membretada'].search(
            [('active', '=', True), ('name', '=', 'Hoja Blanca')], limit=1
        )
        if not hoja:
            hoja = self.env['reclutamiento__kuale.hoja_membretada'].search(
                [('active', '=', True)], limit=1
            )

        return self.create({
            'applicant_id': applicant_id,
            'format_id': format_id,
            'hoja_id': hoja.id if hoja else False,
            'body_html': body,
        })

    def _render_format_body(self, applicant, formato):
        from datetime import date
        import locale
        import re
        from html import unescape

        try:
            template = formato.body or ''
            context = {'object': applicant}

            def eval_expr(expr, ctx):
                expr = expr.strip()
                if expr == "object.actual_date":
                    try:
                        locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')
                    except Exception:
                        try:
                            locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
                        except Exception:
                            pass
                    return date.today().strftime('%d de %B del %Y')
                try:
                    value = eval(expr, dict(ctx))
                    if value is False or value is None:
                        return ''
                    return str(value)
                except Exception:
                    return ''

            def process_tout(template_str, ctx):
                # Maneja <t t-out="..."></t>
                result = re.sub(
                    r'<t\s+t-out="([^"]+)"\s*></t>',
                    lambda m: eval_expr(m.group(1), ctx),
                    template_str
                )
                # Maneja <t t-out="..."/>
                result = re.sub(
                    r'<t\s+t-out="([^"]+)"\s*/>',
                    lambda m: eval_expr(m.group(1), ctx),
                    result
                )
                return result

            def process_foreach(template_str, ctx):
                foreach_pattern = re.compile(
                    r'<t\s+t-foreach="([^"]+)"\s+t-as="([^"]+)"\s*>(.*?)</t>',
                    re.DOTALL
                )

                def replace_foreach(match):
                    collection_expr = match.group(1).strip()
                    var_name = match.group(2).strip()
                    inner_template = match.group(3)
                    try:
                        collection = eval(collection_expr, dict(ctx))
                    except Exception:
                        return ''
                    result = ''
                    for item in collection:
                        item_ctx = dict(ctx)
                        item_ctx[var_name] = item
                        result += process_tout(inner_template, item_ctx)
                    return result

                return foreach_pattern.sub(replace_foreach, template_str)

            template = unescape(str(template))
            # 1. Procesar t-foreach (loops)
            body_rendered = process_foreach(template, context)
            # 2. Procesar t-out restantes (fuera de loops)
            body_rendered = process_tout(body_rendered, context)

            return Markup(body_rendered)

        except Exception as e:
            _logger.error("ERROR _render_format_body: %s", e)
            return Markup('<p>Error al renderizar el formato</p>')

    def action_close(self):
        line = self.env['reclutamiento__kuale.formatos_wizard_line'].search([
            ('format_id', '=', self.format_id.id),
        ], limit=1)
        if line and line.wizard_id:
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'reclutamiento__kuale.formatos_wizard',
                'view_mode': 'form',
                'res_id': line.wizard_id.id,
                'target': 'new',
                'context': {'dialog_size': 'medium'},
            }
        return {'type': 'ir.actions.act_window_close'}

    # ─────────────────────────────────────────────────────────────────────────
    # BACKGROUND IMAGE  (sin cambios, sigue funcionando bien)
    # ─────────────────────────────────────────────────────────────────────────
    def _add_background_image(self, doc, section, img_path):
        """Agrega imagen como fondo de página completa en el header."""
        from lxml import etree
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import tempfile, os

        tmp_path = None
        use_path = img_path
        try:
            from PIL import Image as PILImage
            tmp_path = tempfile.mktemp(suffix='.png')
            img = PILImage.open(img_path)
            img = img.convert('RGB')
            img.save(tmp_path, 'PNG')
            use_path = tmp_path
        except Exception as e:
            _logger.error("Pillow error: %s", e)

        header = section.header
        header.is_linked_to_previous = False

        rId, _ = header.part.get_or_add_image(use_path)

        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

        # Dimensiones carta en EMUs (English Metric Units)
        page_w = int(8.5 * 914400)
        page_h = int(11 * 914400)

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()

        # Párrafo con altura CERO para que no consuma espacio visible
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '20')
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)

        run = para.add_run()

        # Run de 1pt para no agregar altura
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '2')
        rPr.append(sz)
        run._r.insert(0, rPr)

        xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
            xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
        <w:drawing>
            <wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0"
                    simplePos="0" relativeHeight="251658240" locked="0"
                    layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
            <wp:extent cx="{page_w}" cy="{page_h}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="1" name="Background"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
                <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic>
                    <pic:nvPicPr>
                    <pic:cNvPr id="1" name="Background"/>
                    <pic:cNvPicPr/>
                    </pic:nvPicPr>
                    <pic:blipFill>
                    <a:blip r:embed="{rId}"/>
                    <a:stretch><a:fillRect/></a:stretch>
                    </pic:blipFill>
                    <pic:spPr>
                    <a:xfrm>
                        <a:off x="0" y="0"/>
                        <a:ext cx="{page_w}" cy="{page_h}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    </pic:spPr>
                </pic:pic>
                </a:graphicData>
            </a:graphic>
            </wp:anchor>
        </w:drawing>
        </w:r>'''

        run._r.addprevious(etree.fromstring(xml))

    # ─────────────────────────────────────────────────────────────────────────
    # HTML → DOCX  (reescrito con BeautifulSoup)
    # ─────────────────────────────────────────────────────────────────────────
    def _add_html_content_to_doc(self, doc, html):
        """
        Convierte HTML a contenido docx preservando:
        - Alineación de párrafos (left, center, right, justify)
        - Negritas, cursiva, subrayado
        - Tablas (con bordes)
        - Listas ul / ol
        - Saltos de línea <br>
        """
        try:
            from bs4 import BeautifulSoup, NavigableString, Tag
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            if not html:
                return

            soup = BeautifulSoup(html, 'html.parser')

            ALIGN_MAP = {
                'center':  WD_ALIGN_PARAGRAPH.CENTER,
                'right':   WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'left':    WD_ALIGN_PARAGRAPH.LEFT,
            }
            FONT_NAME = 'Arial'
            FONT_SIZE = Pt(10.5)
            LINE_SP   = Pt(13)

            # ── helpers ──────────────────────────────────────────────────────

            def _get_align(el):
                """Lee text-align del style o del atributo align."""
                if not isinstance(el, Tag):
                    return 'justify'
                style = el.get('style', '')
                m = re.search(r'text-align\s*:\s*([\w-]+)', style, re.I)
                if m:
                    return m.group(1).lower()
                return el.get('align', 'justify').lower()

            def _set_para_fmt(p, align='justify', empty=False):
                p.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6) if empty else Pt(0)
                if not empty:
                    p.paragraph_format.line_spacing = LINE_SP

            def _add_run(p, text, bold=False, italic=False, underline=False):
                if not text:
                    return
                run = p.add_run(text)
                run.font.name    = FONT_NAME
                run.font.size    = FONT_SIZE
                run.bold         = bold      or None
                run.italic       = italic    or None
                run.underline    = underline or None

            def _inline(p, node, bold=False, italic=False, underline=False):
                """Recorre nodos inline y agrega runs al párrafo."""
                if isinstance(node, NavigableString):
                    text = str(node)
                    # Preservar espacios internos pero no saltos de línea solos
                    text = text.replace('\n', ' ')
                    if text:
                        _add_run(p, text, bold, italic, underline)
                elif isinstance(node, Tag):
                    if node.name in ('strong', 'b'):
                        for c in node.children:
                            _inline(p, c, True, italic, underline)
                    elif node.name in ('em', 'i'):
                        for c in node.children:
                            _inline(p, c, bold, True, underline)
                    elif node.name == 'u':
                        for c in node.children:
                            _inline(p, c, bold, italic, True)
                    elif node.name == 'br':
                        run = p.add_run()
                        run.add_break()
                    elif node.name == 'span':
                        # Heredar bold de style font-weight
                        style = node.get('style', '')
                        is_bold = bold or bool(re.search(r'font-weight\s*:\s*(bold|[6-9]\d\d)', style, re.I))
                        for c in node.children:
                            _inline(p, c, is_bold, italic, underline)
                    else:
                        for c in node.children:
                            _inline(p, c, bold, italic, underline)

            # ── bloque principal ─────────────────────────────────────────────

            BLOCK_TAGS = {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                          'table', 'ul', 'ol', 'hr'}

            def _process_block(el):
                if isinstance(el, NavigableString):
                    text = str(el).strip()
                    if text:
                        p = doc.add_paragraph()
                        _set_para_fmt(p)
                        _add_run(p, text)
                    return

                if not isinstance(el, Tag):
                    return

                name = el.name

                # ── tabla ────────────────────────────────────────────────────
                if name == 'table':
                    _process_table(el)

                # ── listas ───────────────────────────────────────────────────
                elif name in ('ul', 'ol'):
                    style = 'List Bullet' if name == 'ul' else 'List Number'
                    for li in el.find_all('li', recursive=False):
                        p = doc.add_paragraph(style=style)
                        _set_para_fmt(p)
                        for c in li.children:
                            _inline(p, c)

                # ── separador ───────────────────────────────────────────────
                elif name == 'hr':
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)

                # ── salto de línea suelto ────────────────────────────────────
                elif name == 'br':
                    p = doc.add_paragraph()
                    _set_para_fmt(p, empty=True)

                # ── párrafos / divs / headings ───────────────────────────────
                elif name in ('p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    # ¿Contiene sub-bloques?
                    has_sub_blocks = any(
                        isinstance(c, Tag) and c.name in BLOCK_TAGS
                        for c in el.children
                    )
                    if has_sub_blocks:
                        for c in el.children:
                            _process_block(c)
                    else:
                        align = _get_align(el)
                        p = doc.add_paragraph()
                        # Headings en negritas automáticas
                        auto_bold = name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6')
                        _set_para_fmt(p, align)
                        for c in el.children:
                            _inline(p, c, bold=auto_bold)
                        # Si el párrafo quedó vacío, marcar como separador
                        if not p.text.strip():
                            _set_para_fmt(p, align, empty=True)

            # ── tabla helper ─────────────────────────────────────────────────
            def _process_table(table_el):
                rows = table_el.find_all('tr')
                if not rows:
                    return
                max_cols = max(
                    len(row.find_all(['td', 'th'])) for row in rows
                ) if rows else 0
                if max_cols == 0:
                    return

                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j >= max_cols:
                            break
                        tc = tbl.cell(i, j)
                        # Limpiar contenido por defecto
                        for old_p in tc.paragraphs:
                            old_p.clear()
                        p = tc.paragraphs[0]
                        align = _get_align(cell)
                        _set_para_fmt(p, align)
                        is_header = cell.name == 'th'
                        for c in cell.children:
                            _inline(p, c, bold=is_header)
                        # Asegurar fuente en todos los runs
                        for run in p.runs:
                            if not run.font.name:
                                run.font.name = FONT_NAME
                            if not run.font.size:
                                run.font.size = FONT_SIZE

            # ── iterar hijos del documento ───────────────────────────────────
            for child in soup.children:
                _process_block(child)

        except Exception:
            import traceback
            _logger.error("Error _add_html_content_to_doc:\n%s", traceback.format_exc())
            # Fallback: texto plano
            from bs4 import BeautifulSoup
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            soup2 = BeautifulSoup(html, 'html.parser')
            for line in soup2.get_text('\n').splitlines():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.strip())
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)

    # ─────────────────────────────────────────────────────────────────────────
    # DESCARGA WORD
    # ─────────────────────────────────────────────────────────────────────────
    def action_download_word(self):
        """Genera .docx con membretada de fondo y lo descarga."""
        self.ensure_one()
        try:
            import os, base64, io
            from docx import Document
            from docx.shared import Pt, Cm

            addon_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            img_filename = self.hoja_id.filename if self.hoja_id and self.hoja_id.filename else ''
            img_path = os.path.join(
                addon_path, 'static', 'src', 'hojas_membretadas', img_filename
            ) if img_filename else ''

            doc = Document()
            section = doc.sections[0]

            # ── Tamaño carta ─────────────────────────────────────────────────
            section.page_height = Cm(27.94)
            section.page_width  = Cm(21.59)

            # ── Márgenes normales ─────────────────────────────────────────────
            PAGE_H_CM = 27.94
            PAGE_W_CM = 21.59
            PREVIEW_H_PX = 1056
            PREVIEW_W_PX = 816

            top_px    = 220
            bottom_px = 150
            left_px   = 75
            right_px  = 75

            section.top_margin      = Cm(round(PAGE_H_CM * top_px    / PREVIEW_H_PX, 2))
            section.bottom_margin   = Cm(round(PAGE_H_CM * bottom_px / PREVIEW_H_PX, 2))
            section.left_margin     = Cm(round(PAGE_W_CM * left_px   / PREVIEW_W_PX, 2))
            section.right_margin    = Cm(round(PAGE_W_CM * right_px  / PREVIEW_W_PX, 2))
            section.header_distance = Cm(0.5)
            section.footer_distance = Cm(0.5)

            # ── Membretada como fondo ─────────────────────────────────────────
            if img_path and os.path.exists(img_path):
                self._add_background_image(doc, section, img_path)
            else:
                _logger.warning("Imagen membretada no encontrada: %s", img_path)

            # ── Contenido HTML → DOCX ─────────────────────────────────────────
            self._add_html_content_to_doc(doc, str(self.body_html or ''))

            # ── Guardar y devolver ────────────────────────────────────────────
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            applicant_name = (self.applicant_id.partner_name or 'candidato').replace(' ', '_')
            formato_name   = (self.format_id.name or 'formato').replace(' ', '_')
            filename = f"{formato_name}_{applicant_name}.docx"

            attachment = self.env['ir.attachment'].sudo().create({
                'name': filename,
                'type': 'binary',
                'datas': base64.b64encode(buffer.read()),
                'mimetype': (
                    'application/vnd.openxmlformats-officedocument'
                    '.wordprocessingml.document'
                ),
                'res_model': 'reclutamiento__kuale.formatos_preview_wizard',
                'res_id': self.id,
            })

            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'new',
            }

        except Exception:
            import traceback
            tb = traceback.format_exc()
            _logger.error("Error generando Word:\n%s", tb)
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': 'Error',
                    'message': f'No se pudo generar el archivo Word:\n{tb}',
                    'type': 'danger',
                    'sticky': True,
                }
            }

    def action_print(self):
        return self.action_download_word()